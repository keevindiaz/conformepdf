import pythoncom
from flask import Flask, render_template, request, send_file
import os
from docx import Document
import comtypes.client

app = Flask(__name__)

@app.route('/')
def index():
    return render_template('form.html')

@app.route('/generate_pdf', methods=['POST'])
def generate_pdf():
    pythoncom.CoInitialize()
    
    fecha = request.form.get('fecha', '').strip()
    cliente = request.form.get('cliente', '').strip()
    proyecto = request.form.get('proyecto', '').strip()
    trato = request.form.get('trato', '').strip()

    doc_path = os.path.join(os.path.dirname(__file__), 'plantilla.docx')
    output_doc_path = "temp.docx"
    output_pdf_path = f"{cliente} - {proyecto} - {trato}.pdf"

    doc = Document(doc_path)

    def replace_text(doc, old_text, new_text):
        for para in doc.paragraphs:
            for run in para.runs:
                if old_text in run.text:
                    run.text = run.text.replace(old_text, new_text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if old_text in run.text:
                                run.text = run.text.replace(old_text, new_text)
        for section in doc.sections:
            for para in section.header.paragraphs + section.footer.paragraphs:
                for run in para.runs:
                    if old_text in run.text:
                        run.text = run.text.replace(old_text, new_text)

    replace_text(doc, "{{FECHA}}", fecha)
    replace_text(doc, "{{CLIENTE}}", cliente)
    replace_text(doc, "{{PROYECTO}}", proyecto)
    replace_text(doc, "{{TRATO}}", trato)

    doc.save(output_doc_path)

    word = comtypes.client.CreateObject('Word.Application')
    word.Visible = False
    doc = word.Documents.Open(os.path.abspath(output_doc_path))
    doc.SaveAs(os.path.abspath(output_pdf_path), FileFormat=17)
    doc.Close()
    word.Quit()

    return send_file(output_pdf_path, as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True)
